"""Builds every synthetic test file. Deterministic content, no randomness."""
import csv
import datetime
import json
from pathlib import Path

import docx
import openpyxl

OFFICIAL_RULES = [
    {"rule_id": "V-1001", "title": "Password reuse must be restricted",
     "severity": "high",
     "check_text": "Run SHOW PARAMETER password_reuse_max and verify the value.",
     "fix_text": "Set password_reuse_max to 9 or more.",
     "expected_value": "9 or more"},
    {"rule_id": "V-1002", "title": "Maximum password age must be limited",
     "severity": "medium",
     "check_text": "Verify maximum password age is no greater than 60 days.",
     "fix_text": "Set maximum password age to 60 days.",
     "expected_value": "60 days or less"},
    {"rule_id": "V-1003", "title": "Audit logging must be enabled",
     "severity": "high",
     "check_text": "Verify audit logging is enabled for all databases.",
     "fix_text": "Enable audit logging.",
     "expected_value": "enabled"},
    {"rule_id": "V-1004", "title": "Session timeout must be enforced",
     "severity": "medium",
     "check_text": "Verify idle session timeout is 15 minutes or less.",
     "fix_text": "Set session timeout to 15 minutes.",
     "expected_value": "15 minutes or less"},
    {"rule_id": "V-1005", "title": "Minimum password length must be enforced",
     "severity": "high",
     "check_text": "Verify minimum password length is at least 14 characters.",
     "fix_text": "Set minimum password length to 14.",
     "expected_value": "14 or more"},
]

COMPANY_ROWS = [
    ["High", "Password reuse must be restricted",
     "Database users should not reuse recent passwords",
     "Run SHOW PARAMETER password_reuse_max", "9 or more", "9"],
    ["Medium", "Passwords must be rotated at least every 60 days",
     "Password aging is configured on all accounts",
     "Check profile PASSWORD_LIFE_TIME", "60", "60"],
    ["High", "Logging checked", "Logging was reviewed", "", "", ""],
    ["Low", "Screensaver must show corporate logo",
     "Branding requirement from marketing", "Visual inspection", "logo.png", ""],
]

_HEADERS = ["Group", "STIG Requirement", "Description",
            "Command to Verify", "Approved Setting", "Observed Value"]
_MESSY_HEADERS = ["Area", "What we did", "How we checked", "Value"]
_CSV_COLS = ["Rule ID", "Title", "Severity", "Check Text", "Fix Text",
             "Expected Value"]

EX1_HEADERS = ["STIG REQUIREMENT", "DESCRIPTION", "COMMAND TO VERIFY",
               "APPROVED SETTING"]
EX1_ROWS = [
    ["Password reuse must be restricted",
     "Database users should not reuse recent passwords",
     "Run SHOW PARAMETER password_reuse_max", "9 or more"],
    ["Audit logging must be enabled",
     "Audit logging is required for all databases",
     "Verify audit logging is enabled", "enabled"],
]

EX2_HEADERS = ["", "System Value/Parameter", "Description",
               "REPORTING yes/no", "ENFORCING YES/NO",
               "ADOPT COMPANY STANDARDS DEVIATION/COMPLY",
               "COMPANY AGREED SETTING/COMMAND TO IMPLEMENT", "SEVERITY",
               "CURRENT SETTING", "REMARKS/JUSTIFICATION"]
EX2_ROWS = [
    ["1", "Session timeout must be enforced",
     "Idle session timeout is 15 minutes or less", "YES", "YES", "COMPLY",
     "Set session timeout to 15 minutes", "MEDIUM", "15", ""],
    ["2", "Minimum password length must be enforced",
     "Minimum password length is at least 14 characters", "YES", "NO",
     "DEVIATION", "Set minimum password length to 14", "HIGH", "10",
     "Legacy app cannot handle 14 characters"],
]

INSTRUCTIONS_ROWS = [
    ["Step", "Instruction"],
    ["1", "Fill in every table below before submission."],
    ["2", "Email the completed document to the security team."],
]

GENERAL_INFO_ROWS = [
    ["Field", "Value"],
    ["Application name", "Payments Gateway"],
    ["Team", "Platform Engineering"],
]


def _write_official_csv(path, rules):
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(_CSV_COLS)
        for r in rules:
            w.writerow([r["rule_id"], r["title"], r["severity"],
                        r["check_text"], r["fix_text"], r["expected_value"]])


def _write_docx(path, headers, rows):
    d = docx.Document()
    t = d.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row[: len(headers)]):
            cells[i].text = val
    d.save(str(path))


def _write_real_docx(path):
    d = docx.Document()
    sections = [
        ("General Information", GENERAL_INFO_ROWS[0], GENERAL_INFO_ROWS[1:]),
        ("Instructions", INSTRUCTIONS_ROWS[0], INSTRUCTIONS_ROWS[1:]),
        ("JB.1.1 STIG HARDEING- SEVERITY HIGH", EX1_HEADERS, EX1_ROWS),
        ("IM-1.1 Settings related to Policy or Standards", EX2_HEADERS,
         EX2_ROWS),
    ]
    for heading, headers, rows in sections:
        d.add_paragraph(heading)
        t = d.add_table(rows=1, cols=len(headers))
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row[: len(headers)]):
                cells[i].text = val
    d.save(str(path))


def _write_merged_docx(path):
    """Write a docx with a heading and a table with merged cells in the second data row."""
    d = docx.Document()
    d.add_paragraph("Merged Example")
    headers = ["A", "B"]
    t = d.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    # First data row (no merge)
    cells = t.add_row().cells
    cells[0].text = "Cell A1"
    cells[1].text = "Cell B1"
    # Second data row (merged)
    cells = t.add_row().cells
    cells[0].text = "Merged A2"
    cells[1].text = "Merged B2"
    cells[0].merge(cells[1])
    d.save(str(path))


def build_all(out_dir):
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    paths = {}

    paths["official_csv"] = out / "official.csv"
    _write_official_csv(paths["official_csv"], OFFICIAL_RULES)

    paths["official_json"] = out / "official.json"
    paths["official_json"].write_text(
        json.dumps(OFFICIAL_RULES, indent=1), encoding="utf-8")

    paths["official_xlsx"] = out / "official.xlsx"
    wb = openpyxl.Workbook()
    wb.properties.created = datetime.datetime(2024, 1, 1)
    wb.properties.modified = datetime.datetime(2024, 1, 1)
    ws = wb.active
    ws.title = "Rules"
    ws.append(_CSV_COLS)
    for r in OFFICIAL_RULES:
        ws.append([r["rule_id"], r["title"], r["severity"],
                   r["check_text"], r["fix_text"], r["expected_value"]])
    wb.save(str(paths["official_xlsx"]))

    dup = [dict(OFFICIAL_RULES[0], rule_id="V-2001"),
           dict(OFFICIAL_RULES[1], rule_id="V-2001")]
    paths["official_dup_ids_csv"] = out / "official_dup.csv"
    _write_official_csv(paths["official_dup_ids_csv"], dup)

    paths["company_docx"] = out / "company.docx"
    _write_docx(paths["company_docx"], _HEADERS, COMPANY_ROWS)

    messy = [[r[0], r[1] + ". " + r[2], r[3], r[4]] for r in COMPANY_ROWS]
    messy.append(["", "", "", ""])                     # empty row
    paths["company_docx_messy"] = out / "company_messy.docx"
    _write_docx(paths["company_docx_messy"], _MESSY_HEADERS, messy)

    paths["company_xlsx"] = out / "company.xlsx"
    wb = openpyxl.Workbook()
    wb.properties.created = datetime.datetime(2024, 1, 1)
    wb.properties.modified = datetime.datetime(2024, 1, 1)
    ws = wb.active
    ws.title = "Submission"
    ws.append(_HEADERS)
    for row in COMPANY_ROWS:
        ws.append(row)
    wb.save(str(paths["company_xlsx"]))

    paths["company_real_docx"] = out / "company_real.docx"
    _write_real_docx(paths["company_real_docx"])

    paths["company_merged_docx"] = out / "company_merged.docx"
    _write_merged_docx(paths["company_merged_docx"])

    return paths
